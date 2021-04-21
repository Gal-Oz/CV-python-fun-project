from docx import Document    #ספריה ושימוש בפונקציה מסמכים
from docx.shared import Inches #שימוש בפונקציה לשינוי גודל במסמך
import pyttsx3     #שימוש בספריה של טקסט לדיבור

def speak(text):
    pyttsx3.speak(text)

speak('Hi I am artificial intelligence,  I can change at any moment according to your choices'' , ' 'for example my voice.')
engine = pyttsx3.init()    #שינוי קול מכני של אישה לקול של גבר
voices = engine.getProperty('voices')
for voice in voices:
   engine.setProperty('voice', voice.id)
engine.runAndWait()

speak('I am interested in knowing what your name is ')
document = Document()        #הגדרת פונקציה כמשתנה

#תמונת פרופיל
document.add_picture('me.jpg',width=Inches(0.95))

#פרטים אישיים ליצירת קשר
name = input('What is your name? ')
last_name= input('What is your  last name? ')
speak('Hello ' + name +' ,  ' + last_name + '  how are you today?')     #המערכת מברכת את המשתמש על שימוש בתוכנית
speak('I am not trying to flirt with you, What is your phon number?')
phon_N=input('What is your phon number? ')
speak('What is your email address?')
email=input('What is your email address? ')

#בניית תוכן
document.add_paragraph(
name + ' '+ last_name + ' | ' + phon_N + ' | ' + email)    #הוספת פונקציה של שורת תוכן במסמך

#about me
document.add_heading('About me')
speak('tell about yourself: ')
about_me = input('tell about yourself: ')
document.add_paragraph(about_me)

#work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('plese enter company name')
company = input('Enter company: ')
speak('When did you start ')
from_date = input('From Date: ')
speak('How long? ')
to_date = input('To Date: ')
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True
speak('Decribe your experience at ' + company +', ')
experience_details = input('Decribe your experience at ' + company +': ')
p.add_run(experience_details)

#more experiences
while True:
    speak('Do you have more experiences? ')
    has_more_experiences = input('Do you have more experiences? Yes or No:  ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        speak('Amazing ,  plese enter company name')
        company = input('Enter company: ')
        speak('When did you start ')
        from_date = input('From Date: ')
        speak('How long? ')
        to_date = input('To Date: ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        speak('Decribe your experience at ' + company +': ')
        experience_details = input('Decribe your experience at ' + company +': ')
        p.add_run(experience_details)
    else:
        speak('Sucks, but you whill succeed next time')
        break


#skills
document.add_heading('Skills')
speak('You are talented , plese , enter one of your skills ')
skill = input('Enter skills: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('Do you have more skills? ')
    has_more_skills = input('Do you have more skills? Yes or No:  ')
    if has_more_skills.lower() == 'yes':
        speak('Amazing ,  plese enter skill')
        skill = input('Enter skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        speak('Are you sure? You are really talented')
        break

speak('Completed the Curriculum Vitae is ready! ' + 'Good luck with your next job ')
    #footer   כל הזכויות שמורות / אודות
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generat using be PYTHON code"


document.save('cv.docx')  #שמירת המסמך על שם cv